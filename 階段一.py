"""完成版-2.2(pf,eff)"""
import os
import fitz
import pytesseract
import PyPDF2
from PIL import Image
import pandas as pd
import json
import google.generativeai as genai
import time
from docx import Document
from dotenv import load_dotenv
load_dotenv()

# 設定您的 Gemini API 金鑰
# 建議將金鑰存放在環境變數中，以確保安全
genai.configure(api_key=os.getenv("GEMINI_API_KEY"))
#pytesseract.pytesseract.tesseract_cmd = r"C:\Program Files\Tesseract-OCR\tesseract.exe"
# PDF 路徑
#pdf_path = r"D:\someDATA\96-40318-02@A0.pdf"              #最早的HUB250            無法成常抽取
#pdf_path = r"D:\someDATA\20250928\96-40321-01@I0.pdf"     #HUB250-111             可成常抽取
#pdf_path = r"D:\someDATA\20250928\96-42403-01@G0.pdf"     #HPA65X-240             可成常抽取
#pdf_path = r"D:\someDATA\20250928\96-42404-01@F0.pdf"     #HPA65X-120             可成常抽取
#pdf_path = r"D:\someDATA\20250928\96-42998-02@A0.pdf"     #HPA90X-120             可成常抽取
def process_pdf_to_excel(pdf_path: str, output_excel: str):
        # (注意：下一行程式碼就是你原本的 doc = fitz.open(...))
    #doc = fitz.open(pdf_path)

    # 設定 Tesseract 路徑
    
    #設定excel輸出
    #output_excel = r"D:\someDATA\function_test_spec HPA65X-120-3.xlsx"
    #設定Word輸出
    #output_docx = r"D:\someDATA\output.docx"
    # 輸出資料夾 (用於存放暫存圖片)
    ##os.makedirs(output_folder, exist_ok=True)

    # 建立 Word 文件
    #word_doc = Document()

    # 累積所有辨識出來的文字
    full_ocr_text = ""
    print(pdf_path)
    with open(pdf_path, 'rb') as pdfFileObj:
        pdfReader = PyPDF2.PdfReader(pdfFileObj)
        num_pages = len(pdfReader.pages)

        for i in range(num_pages):
            pageObj = pdfReader.pages[i]   # ✅ use pages[i] instead of getPage(i)
            full_ocr_text += pageObj.extract_text() # ✅ method renamed to extract_text()

    if "TEST" in full_ocr_text:
        print("可正常讀取，跳過OCR")
    else:
        with fitz.open(pdf_path) as doc:
            print("無法正常讀取，使用OCR抓取資料")

            for page_num in range(len(doc)):
                print(f"正在處理第 {page_num + 1} 頁...")

                # 取出每一頁轉成圖片
                page = doc[page_num]
                pix = page.get_pixmap(dpi=200)
                img = Image.frombytes("RGB", [pix.width, pix.height], pix.samples)

                # 儲存圖片檔
                #img_path = os.path.join(output_folder, f"page_{page_num+1}.png")
                #img.save(img_path)
            
                # OCR 辨識（英文）
                text = pytesseract.image_to_string(img, lang="eng")
                full_ocr_text += text + "\n"  # 將每頁文字累加並加上換行
            
                # 每一頁存進 Word（加上頁碼標題）
                #word_doc.add_heading(f"Page {page_num + 1}", level=1)
                #word_doc.add_paragraph(text)

            print("OCR完成")
            #word_doc.save(output_docx)
            #print("OCR文字Word檔案輸出")
            #doc.close()

    # --- Gemini API 接入 ---
    print("正在將文字傳送給 Gemini...")

    # 設計給 Gemini 的 Prompt
    prompt = f"""
    你是個能將 OCR 擷取文字轉換成結構化 JSON 格式的專家。
    請從以下「功能測試規格」文件中的表格內容，提取出 "TEST PARAMETER", "TEST CONDITION" I,"TEST CONDITION II","TEST CONDITION III","TEST CONDITION IV", 和 "LIMITS I","LIMITS II","LIMITS III", 的資訊。

    "Dropout Point"並不是"TEST PARAMETER"的內容，不要將"Dropout Point"填在"TEST PARAMETER"，請填在TEST CONDITION IV，"Dropout Point:"文字不出現。
    "LIMITS I"中的參數不一定一樣，請分開寫
    "TEST CONDITION" I,"TEST CONDITION II","TEST CONDITION III","TEST CONDITION IV"分別為Voltage,Frequency,Load,Dropout Point。
    "Vin","VAC","Io1","Io1,2","test 2s"文字不要出現

    TEST CONDITION I出現的電壓不用標"V"

    輸出中出現多組電壓/頻率時 輸出為多行並且TEST PARAMETER相同時item相同
    PEAK LOAD TEST是例外，遇上多組電壓時只輸出小的項目

    Load 出現以\分開的多組內容時分開多行並且TEST PARAMETER相同時item相同輸出，同時電壓有多種時每種電壓都要有多種Load，而以,或其他'方式分開時不分行
    Dynamic Lood test有多種負載時不用分開測

    填入時請將"Tin"更改為"Iin"，"VAC"前面不應出現"O"，請訂正成"0VAC"。
    在"Vo1"後的"V"放"LIMITS I","mVpp"放"LIMITS II"。
    在"Vo2"後面的參數例如"Vo2: +5.00 to +8.00 V"放置於"LIMITS III"。
    "LIMITS I"與"LIMITS II"的參數不會出現"A"，假如參數有"A"則放到"LIMITS IV"去。
    電流Iin的參數放"LIMITS IV",Pin:參數放PIN,如果LIMITS出現XX%，XX填入 "LIMITS VI"，出現PF參數填入"LIMITS VII"例如"Iin:1.16A max"要填在"LIMITS IV"，剩下未歸類的放"LIMITS VIII"
    "TS: 1S,20mS"等放在"LIMITS V"，"TS:"文字不顯示。
    "LINE AND LOAD REGULATION TEST"只要"LIMITS IV"有Iin數值"LIMITS II"就會是空的

    TEST PARAMETER類別的文字請使用原始英文寫法，例如SHORT CKT TEST,EFFICIENCY,LINE AND LOAD REGULATION TEST#1,LINE AND LOAD REGULATION TEST#2等
    當文字中出現"as item "時請提取對應item的"LIMITS I"與"LIMITS II"值並填入該項"LIMITS I"與"LIMITS II"例如假設"LINE AND LOAD REGULATION TEST#8" 與item 11與"POWER CONSUMPTION TEST"讀出 as item 6 ，在item 6 "LIMITS I"為 +10.63 to +11.47 V,"LIMITS II"為50mVpp max，則在"LINE AND LOAD REGULATION TEST#8"與item 11 與"POWER CONSUMPTION TEST"的"LIMITS I"填入 +10.63 to +11.47 V,"LIMITS II" 填入50mVpp max依此類推。
    注意上一條只是舉例，請不要將"+10.63 to +11.47 V","50mVpp max"填入。.

    將每個測試項目作為一個 JSON 物件，並以一個完整的 JSON 陣列返回。
    JSON 陣列中的每個物件必須包含以下十一個鍵名 (key)："item,"TEST PARAMETER", "TEST CONDITION" I,"TEST CONDITION II","TEST CONDITION III","TEST CONDITION IV", 和 "LIMITS I","LIMITS II","LIMITS III","LIMITS IV,"LIMITS V","PIN","LIMITS VI","LIMITS VII","LIMITS VIII。

    OCR 擷取文字如下：
    {full_ocr_text}
    """

    # 加入重試機制
    max_retries = 3
    for attempt in range(max_retries):
        try:
            model = genai.GenerativeModel('gemini-2.5-flash')
            response = model.generate_content(prompt)
            if not response or not response.text:
                raise Exception("API 未返回有效內容。")
            # 1. 取得原始文字
            raw_json_text = response.text
            # 2. 清理開頭和結尾的 Markdown 標記
            if raw_json_text.startswith('```json'):
                raw_json_text = raw_json_text[len('```json'):].strip()
            if raw_json_text.endswith('```'):
                raw_json_text = raw_json_text[:-len('```')].strip()
            # 3. 解析清理後的 JSON 字串
            data_list = json.loads(raw_json_text)
            # 建立 pandas DataFrame
            df = pd.DataFrame(data_list)
            # 將 DataFrame 存成 Excel 檔案
            df.to_excel(output_excel, index=False)

            print("-" * 30)
            print(f"✅ 成功！試算表已成功生成：{output_excel}")
            return True
        except json.JSONDecodeError as e:
            print("❌ 錯誤：Gemini 返回的內容不是有效的 JSON 格式。")
            print(f"JSON 解析錯誤：{e}")
            print("Gemini 返回的原始內容：\n", response.text)
            return False

        except Exception as e:
            if "429" in str(e):
                print(f"❌ 發生額度超限錯誤，正在等待 {60 * (attempt + 1)} 秒後重試... (第 {attempt + 1} 次)")
                time.sleep(60 * (attempt + 1))
            else:
                print(f"❌ 發生其他錯誤：{e}")
                return False
              

    else:
        print("❌ 重試次數已達上限，無法完成請求。請檢查您的 API 額度。")
        return False
    

